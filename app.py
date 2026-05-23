import logging
import json
import os
import pickle
import re
import uuid
from datetime import datetime
from typing import List, Dict, Any, Tuple

import faiss
import numpy as np
import ollama
import streamlit as st
import torch
from docx import Document
from io import BytesIO
from xml.etree import ElementTree as etree
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer, CrossEncoder
from tokenizers import ByteLevelBPETokenizer
from model.config import Config
from model.checkpoint import load_model_checkpoint
from model.transformer import GPT
from collections import Counter
import math

INDEX_FILE = 'data/vector.index'
META_FILE = 'data/chunks.pkl'
DOCS_ROOT = 'docs'
UPLOAD_DIR = os.path.join(DOCS_ROOT, 'uploads')
DEVICE = 'cuda' if torch.cuda.is_available() else 'cpu'
CACHE_FILE = 'data/response_cache.pkl'
FEEDBACK_FILE = 'data/feedback.jsonl'

os.makedirs('data', exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)
logging.basicConfig(level=logging.INFO)

STOPWORDS = {'what','is','the','a','an','how','explain','tell','me','about','can','you','please','to','of','in','on','at','for','with','from','by','and','or'}
DOMAIN_HINTS = {
    'edi': ['edi','850','810','856','997','as2','x12','edifact'],
    'oracle': ['soa','oic','osb','bpel','mediator','oracle','dehydration'],
    'xslt': ['xslt','xpath','xml','xsd','namespace'],
    'seeburger': ['seeburger','bis'],
    'ai': ['ai','transformer','llm','machine','learning']
}
AVAILABLE_MODELS = ['MapMindGPT-Custom','qwen2.5:7b-instruct','llama3:8b','mistral:7b','deepseek-coder:6.7b','codellama:7b']
AVAILABLE_MODES = ['General Chat','Code Assistant','EDI Expert','Oracle Expert','XSLT/XPath Expert','SEEBURGER Expert']

st.set_page_config(page_title='MapMindGPT', layout='wide')


def get_system_prompt(mode):
    prompts = {
        'General Chat': 'You are MapMindGPT. Use provided context when relevant. Be direct and factual.',
        'Code Assistant': 'You are an expert software engineering assistant. Generate clean production-ready code with minimal explanation.',
        'EDI Expert': 'You are an EDI integration expert specializing in ANSI X12, EDIFACT, AS2 and acknowledgements. Be precise.',
        'Oracle Expert': 'You are an Oracle middleware expert for SOA, OIC, OSB, BPEL, dehydration and fault handling. Focus on practical solutions.',
        'XSLT/XPath Expert': 'You are an XSLT/XPath expert. Prefer XSLT 1.0 unless requested otherwise. Provide working code.',
        'SEEBURGER Expert': 'You are a SEEBURGER BIS expert for partner setup, routing and mappings. Give step-by-step guidance.'
    }
    return prompts.get(mode, prompts['General Chat'])


@st.cache_resource
def load_embed_model():
    return SentenceTransformer('all-MiniLM-L6-v2')


@st.cache_resource
def load_reranker():
    try:
        return CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')
    except Exception as e:
        logging.warning(f'Failed to load reranker: {e}')
        return None


@st.cache_resource
def load_custom_model():
    try:
        vocab_path = 'tokenizer/vocab.json'
        merges_path = 'tokenizer/merges.txt'
        checkpoint_path = 'checkpoints/model.pt'

        missing = [path for path in [vocab_path, merges_path, checkpoint_path] if not os.path.exists(path)]
        if missing:
            raise FileNotFoundError(f"Missing custom model files: {', '.join(missing)}")

        tokenizer = ByteLevelBPETokenizer(
            vocab_path,
            merges_path
        )

        model = GPT(Config()).to(DEVICE)
        load_model_checkpoint(model, checkpoint_path, DEVICE)
        model.eval()

        return tokenizer, model

    except Exception as e:
        logging.exception("Custom model load failed")
        st.error(f"Custom model load failed: {e}")
        return None, None


def generate_custom_response(prompt, max_tokens=220):
    tokenizer, model = load_custom_model()
    if tokenizer is None or model is None:
        return 'Custom model not available. Please select an Ollama model.'
    encoded = tokenizer.encode(prompt)
    idx = torch.tensor([encoded.ids], dtype=torch.long).to(DEVICE)
    eos_token_id = tokenizer.token_to_id('<eos>')

    with torch.no_grad():
        output = model.generate(
            idx,
            max_new_tokens=max_tokens,
            temperature=0.75,
            top_k=40,
            top_p=0.9,
            repetition_penalty=1.12,
            eos_token_id=eos_token_id
        )

    new_tokens = output[0, idx.size(1):].tolist()
    response = tokenizer.decode(new_tokens).strip()

    for marker in ['User:', 'Question:', 'Context:', '<eos>']:
        if marker in response:
            response = response.split(marker, 1)[0].strip()

    return response or "I don't have enough signal in the local custom model to answer that well yet."


def is_low_quality_custom_response(response: str) -> bool:
    text = response.strip()
    if len(text) < 20:
        return True

    words = re.findall(r"[A-Za-z][A-Za-z0-9_-]*", text.lower())
    if len(words) < 5:
        return True

    unique_ratio = len(set(words)) / max(len(words), 1)
    repeated_word_count = Counter(words).most_common(1)[0][1]
    symbol_ratio = len(re.findall(r"[^A-Za-z0-9\s.,;:!?()'\"/-]", text)) / max(len(text), 1)
    arrow_noise = text.count("->") + text.count("→")

    return (
        unique_ratio < 0.35
        or repeated_word_count >= 8
        or symbol_ratio > 0.08
        or arrow_noise >= 4
    )


def build_custom_fallback_answer(query: str, context: List[Dict]) -> str:
    keywords = set(extract_keywords(query))
    sentences = []
    detected_domain = detect_domain(query)

    for item in context:
        source = item.get('source', 'source')
        category = item.get('category', 'unknown')
        for sentence in re.split(r"(?<=[.!?])\s+", item.get('text', '')):
            clean = ' '.join(sentence.split())
            if not clean:
                continue
            tokens = set(extract_keywords(clean))
            if keywords and not keywords.intersection(tokens):
                continue
            score = len(keywords.intersection(tokens))
            if detected_domain and category == detected_domain:
                score += 5
            if detected_domain and detected_domain in source.lower():
                score += 2
            if detected_domain == 'edi' and category in {'xslt', 'ai'}:
                score -= 4
            sentences.append((score, clean, source))

    if 'edi' in keywords:
        intro = (
            "EDI, or Electronic Data Interchange, is the structured computer-to-computer "
            "exchange of business documents between trading partners."
        )
        overview = [
            "- It replaces manual exchange of business documents with standardized electronic messages.",
            "- Common EDI documents include purchase orders, invoices, shipment notices, and acknowledgements.",
            "- In ANSI X12, examples include 850 Purchase Order, 810 Invoice, 856 Advance Ship Notice, and 997 Functional Acknowledgement."
        ]
    else:
        intro = "The local custom model could not generate a clean answer, so here is the best answer from the retrieved knowledge base."
        overview = []

    sentences.sort(key=lambda item: item[0], reverse=True)
    selected = [(sentence, source) for _, sentence, source in sentences[:3]]
    if not selected:
        selected = [(format_snippet(item.get('text', ''), length=220), item.get('source', 'source')) for item in context[:2]]

    details = []
    for sentence, source in selected:
        if sentence:
            details.append(f"- {sentence} [{source}]")

    return '\n'.join([intro, *overview, *details])


def is_custom_unavailable_response(response: str) -> bool:
    return response.startswith('Custom model not available')


def read_uploaded_file(uploaded_file):
    try:
        name = uploaded_file.name.lower()
        uploaded_file.seek(0)
        if name.endswith('.txt'):
            return uploaded_file.read().decode('utf-8', errors='ignore')
        if name.endswith('.pdf'):
            reader = PdfReader(BytesIO(uploaded_file.read()))
            return '\n'.join([p.extract_text() for p in reader.pages if p.extract_text()])
        if name.endswith('.docx'):
            doc = Document(uploaded_file)
            return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        if name.endswith('.xml'):
            tree = etree.parse(uploaded_file)
            root = tree.getroot()
            return etree.tostring(root, encoding='unicode')
    except Exception as e:
        logging.exception('Failed to read uploaded file: %s', uploaded_file.name)
        st.error(f'File read failed: {e}')
    return ''


def read_file_from_disk(path):
    try:
        lower = path.lower()
        if lower.endswith('.txt'):
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        if lower.endswith('.pdf'):
            reader = PdfReader(path)
            return '\n'.join([p.extract_text() for p in reader.pages if p.extract_text()])
        if lower.endswith('.docx'):
            doc = Document(path)
            return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        if lower.endswith('.xml'):
            tree = etree.parse(path)
            root = tree.getroot()
            return etree.tostring(root, encoding='unicode')
    except Exception:
        logging.exception('Failed to read file from disk: %s', path)
        return ''
    return ''


def sanitize_filename(name):
    name = os.path.basename(name)
    safe = re.sub(r'[^a-zA-Z0-9_.-]+', '_', name)
    return safe[:200]


def is_uploaded_source(source):
    return source.startswith('uploads/') or source.startswith('uploads\\')


def format_snippet(text, length=260):
    snippet = ' '.join(text.strip().split())
    return snippet[:length] + ('...' if len(snippet) > length else '')


def get_embedding_dim(embed_model):
    if hasattr(embed_model, 'get_sentence_embedding_dimension'):
        return embed_model.get_sentence_embedding_dimension()
    if hasattr(embed_model, 'dimension'):
        return embed_model.dimension
    return 384


def split_text(text, chunk_size=400, overlap=100):
    out = []
    start = 0
    while start < len(text):
        chunk = text[start:start+chunk_size]
        if chunk.strip():
            out.append(chunk)
        start += chunk_size - overlap
    return out


def scan_all_documents():
    docs = []
    for root, _, files in os.walk(DOCS_ROOT):
        for file in files:
            if file.lower().endswith(('.txt','.pdf','.docx','.xml')):
                docs.append(os.path.join(root, file))
    return docs


def rebuild_index():
    embed_model = load_embed_model()
    dim = get_embedding_dim(embed_model)
    all_chunks = []
    for path in scan_all_documents():
        text = read_file_from_disk(path)
        if not text.strip():
            continue
        rel = os.path.relpath(path, DOCS_ROOT)
        category = os.path.basename(os.path.dirname(path)).lower()
        for chunk in split_text(text):
            all_chunks.append({'id': len(all_chunks), 'source': rel, 'category': category, 'text': chunk})
    index = faiss.IndexIDMap(faiss.IndexFlatIP(dim))
    if all_chunks:
        embeddings = embed_model.encode([c['text'] for c in all_chunks], convert_to_numpy=True, normalize_embeddings=True)
        ids = np.arange(len(all_chunks), dtype=np.int64)
        index.add_with_ids(embeddings.astype(np.float32), ids)
    faiss.write_index(index, INDEX_FILE)
    with open(META_FILE, 'wb') as f:
        pickle.dump(all_chunks, f)
    load_resources.clear()
    logging.info('Rebuilt FAISS index with %d chunks and dimension %d', len(all_chunks), dim)


@st.cache_resource
def load_resources():
    if not os.path.exists(INDEX_FILE) or not os.path.exists(META_FILE):
        rebuild_index()
    embed_model = load_embed_model()
    index = faiss.read_index(INDEX_FILE)
    with open(META_FILE, 'rb') as f:
        chunks = pickle.load(f)
    for idx, chunk in enumerate(chunks):
        chunk.setdefault('id', idx)
        chunk.setdefault('source', f'document_{idx}')
        chunk.setdefault('category', 'unknown')
        chunk.setdefault('text', '')
    return embed_model, index, chunks


def add_uploaded_document(uploaded_file):
    text = read_uploaded_file(uploaded_file)
    if not text.strip():
        return False
    filename = sanitize_filename(uploaded_file.name)
    dest_path = os.path.join(UPLOAD_DIR, filename)
    if os.path.exists(dest_path):
        base, ext = os.path.splitext(filename)
        filename = f'{base}_{uuid.uuid4().hex[:8]}{ext}'
        dest_path = os.path.join(UPLOAD_DIR, filename)
    with open(dest_path, 'w', encoding='utf-8') as f:
        f.write(text)
    rebuild_index()
    return True


def get_uploaded_documents():
    return sorted([f for f in os.listdir(UPLOAD_DIR) if os.path.isfile(os.path.join(UPLOAD_DIR, f))])


def delete_uploaded_document(filename):
    path = os.path.join(UPLOAD_DIR, filename)
    if os.path.exists(path):
        os.remove(path)
    rebuild_index()


def clear_uploaded_knowledge():
    for file in get_uploaded_documents():
        os.remove(os.path.join(UPLOAD_DIR, file))
    rebuild_index()


def extract_keywords(query):
    words = re.findall(r'\w+', query.lower())
    return [w for w in words if w not in STOPWORDS]


def detect_domain(query):
    q = query.lower()
    for domain, terms in DOMAIN_HINTS.items():
        if any(term in q for term in terms):
            return domain
    return None


def compute_bm25_scores(query: str, chunks: List[Dict], k1: float = 1.5, b: float = 0.75) -> List[float]:
    keywords = extract_keywords(query)
    if not keywords:
        return [0.0] * len(chunks)
    
    doc_lens = [len(c['text'].split()) for c in chunks]
    avg_len = sum(doc_lens) / len(doc_lens) if doc_lens else 1
    
    df = Counter()
    for c in chunks:
        tokens = set(extract_keywords(c['text']))
        for kw in keywords:
            if kw in tokens:
                df[kw] += 1
    
    N = len(chunks)
    scores = []
    for i, c in enumerate(chunks):
        tokens = extract_keywords(c['text'])
        token_freq = Counter(tokens)
        score = 0.0
        for kw in keywords:
            if kw not in token_freq:
                continue
            tf = token_freq[kw]
            idf = math.log((N - df[kw] + 0.5) / (df[kw] + 0.5) + 1)
            dl = doc_lens[i]
            norm = tf * (k1 + 1) / (tf + k1 * (1 - b + b * dl / avg_len))
            score += idf * norm
        scores.append(score)
    return scores


def retrieve_semantic(query: str, top_k: int, scope: str) -> List[Tuple[float, Dict]]:
    embed_model, index, chunks = load_resources()
    if not chunks:
        return []
    q_emb = embed_model.encode([query], convert_to_numpy=True, normalize_embeddings=True)
    k = min(200, len(chunks))
    distances, ids = index.search(q_emb.astype(np.float32), k)
    
    results = []
    for dist, idx in zip(distances[0], ids[0]):
        if idx < 0:
            continue
        item = chunks[int(idx)]
        if scope == 'Uploaded only' and not is_uploaded_source(item['source']):
            continue
        if scope == 'Built-in docs' and is_uploaded_source(item['source']):
            continue
        score = float(dist)
        results.append((score, item))
    return results[:top_k * 3]


def retrieve_keyword(query: str, top_k: int, scope: str) -> List[Tuple[float, Dict]]:
    _, _, chunks = load_resources()
    if not chunks:
        return []
    
    filtered = chunks
    if scope == 'Uploaded only':
        filtered = [c for c in chunks if is_uploaded_source(c['source'])]
    elif scope == 'Built-in docs':
        filtered = [c for c in chunks if not is_uploaded_source(c['source'])]
    
    if not filtered:
        return []
    
    bm25_scores = compute_bm25_scores(query, filtered)
    results = [(score, chunk) for score, chunk in zip(bm25_scores, filtered)]
    results.sort(key=lambda x: x[0], reverse=True)
    return results[:top_k * 3]


def reciprocal_rank_fusion(semantic: List[Tuple[float, Dict]], keyword: List[Tuple[float, Dict]], k: int = 60) -> List[Dict]:
    scores = {}
    
    for rank, (_, item) in enumerate(semantic):
        key = item['id']
        scores[key] = scores.get(key, 0) + 1.0 / (k + rank + 1)
    
    for rank, (_, item) in enumerate(keyword):
        key = item['id']
        scores[key] = scores.get(key, 0) + 1.0 / (k + rank + 1)
    
    all_items = {item['id']: item for _, item in semantic + keyword}
    ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    return [all_items[item_id] for item_id, _ in ranked]


def rerank_results(query: str, candidates: List[Dict], top_k: int) -> List[Dict]:
    if len(candidates) <= top_k:
        return candidates
    
    reranker = load_reranker()
    if reranker is None:
        return candidates[:top_k]
    
    pairs = [[query, c['text']] for c in candidates]
    scores = reranker.predict(pairs)
    
    ranked = sorted(zip(scores, candidates), key=lambda x: x[0], reverse=True)
    return [item for _, item in ranked[:top_k]]


def retrieve(query: str, top_k: int = 3, scope: str = 'All documents', use_reranker: bool = True) -> List[Dict]:
    semantic = retrieve_semantic(query, top_k, scope)
    keyword = retrieve_keyword(query, top_k, scope)
    
    merged = reciprocal_rank_fusion(semantic, keyword)[:top_k * 3]
    
    if use_reranker and merged:
        return rerank_results(query, merged, top_k)
    return merged[:top_k]


def count_tokens_approx(text: str) -> int:
    return int(len(text.split()) * 1.3)


def truncate_context(context: List[Dict], max_tokens: int = 2000) -> List[Dict]:
    total = 0
    result = []
    for item in context:
        tokens = count_tokens_approx(item['text'])
        if total + tokens > max_tokens:
            break
        result.append(item)
        total += tokens
    return result


def load_cache():
    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, 'rb') as f:
                return pickle.load(f)
        except Exception:
            return {}
    return {}


def save_cache(cache):
    try:
        with open(CACHE_FILE, 'wb') as f:
            pickle.dump(cache, f)
    except Exception as e:
        logging.warning(f'Failed to save cache: {e}')


def cache_key(query, context, model, mode):
    context_ids = tuple(sorted([c.get('id', idx) for idx, c in enumerate(context)]))
    return (query.strip().lower(), context_ids, model, mode)


def ask_llm(query: str, context, history, model_name: str, mode: str, use_cache: bool = True):
    cache = load_cache() if use_cache else {}

    key = cache_key(query, context, model_name, mode)

    if key in cache:
        return cache[key]

    simple_query = query.strip().lower()

    # direct greetings
    if simple_query in ['hi', 'hello', 'hey', 'good morning', 'good evening']:
        return "Hello! How can I help you today?"

    # custom model safety
    if model_name == 'MapMindGPT-Custom':
        if not context:
            return "I don't have relevant knowledge for that query."

        context = truncate_context(context, max_tokens=1200)

        context_text = '\n\n'.join(
            [f"[{c['source']}]\n{c['text']}" for c in context]
        )

        system_prompt = get_system_prompt(mode)

        prompt = f"""
{system_prompt}

Use only the context below. If the context is not enough, say what is missing.
Keep the answer concise and practical.

Context:
{context_text}

User: {query}
Assistant:
"""

        response = generate_custom_response(prompt, max_tokens=180)
        if is_custom_unavailable_response(response) or is_low_quality_custom_response(response):
            response = build_custom_fallback_answer(query, context)

        if use_cache:
            cache[key] = response
            save_cache(cache)

        return response

    # Ollama models
    context = truncate_context(context, max_tokens=2000)

    context_text = '\n\n'.join(
        [f"[{c['source']}]\n{c['text']}" for c in context]
    )

    system_prompt = get_system_prompt(mode)

    messages = [
        {
            'role': 'system',
            'content': system_prompt
        }
    ]

    messages.extend(history[-6:])

    messages.append({
        'role': 'user',
        'content': f'''
Context:
{context_text}

Question:
{query}
'''
    })

    try:
        response_text = ""

        stream = ollama.chat(
            model=model_name,
            messages=messages,
            stream=True
        )

        placeholder = st.empty()

        for chunk in stream:
            if 'message' in chunk and 'content' in chunk['message']:
                response_text += chunk['message']['content']
                placeholder.markdown(response_text)

        if use_cache:
            cache[key] = response_text
            save_cache(cache)

        return response_text

    except Exception as e:
        logging.exception("Ollama failed")
        st.error(f"Model error: {e}")
        return "Error generating response."


def save_feedback(query: str, response: str, rating: str):
    try:
        with open(FEEDBACK_FILE, 'a', encoding='utf-8') as f:
            entry = {'timestamp': datetime.now().isoformat(), 'query': query, 'response': response, 'rating': rating}
            f.write(json.dumps(entry, ensure_ascii=False) + '\n')
    except Exception as e:
        logging.warning(f'Failed to save feedback: {e}')


def export_conversation(messages: List[Dict], format_type: str = 'markdown') -> str:
    if format_type == 'markdown':
        lines = ['# MapMindGPT Conversation', f'\nExported: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n']
        for msg in messages:
            role = 'User' if msg['role'] == 'user' else 'Assistant'
            lines.append(f'\n## {role}\n\n{msg["content"]}\n')
        return '\n'.join(lines)
    elif format_type == 'json':
        import json
        return json.dumps({'timestamp': datetime.now().isoformat(), 'messages': messages}, indent=2)
    return ''


st.title('MapMindGPT')

with st.sidebar:
    st.header('AI Settings')
    selected_model = st.selectbox('Choose Model', AVAILABLE_MODELS)
    selected_mode = st.selectbox('Choose Mode', AVAILABLE_MODES)
    
    st.header('Retrieval Settings')
    use_reranker = st.checkbox('Use reranker', value=True)
    use_cache = st.checkbox('Cache responses', value=True)
    
    embed_model, index, chunks = load_resources()
    index_size = os.path.getsize(INDEX_FILE) if os.path.exists(INDEX_FILE) else 0
    last_rebuild = datetime.fromtimestamp(os.path.getmtime(META_FILE)).strftime('%Y-%m-%d %H:%M:%S') if os.path.exists(META_FILE) else 'n/a'
    st.metric('Indexed chunks', len(chunks))
    st.metric('Index size', f'{index_size // 1024} KB')
    st.write(f'Last rebuild: {last_rebuild}')
    
    st.header('Knowledge Base')
    search_scope = st.radio('Search scope', ['All documents', 'Uploaded only', 'Built-in docs'])
    num_sources = st.slider('Sources to retrieve', 1, 10, 3)
    
    uploaded_file = st.file_uploader('Upload Document', type=['txt','pdf','docx','xml'])
    
    if uploaded_file:
        with st.spinner('Processing document...'):
            if add_uploaded_document(uploaded_file):
                st.success('Document uploaded and indexed.')
                st.rerun()
    
    uploaded_docs = get_uploaded_documents()
    if uploaded_docs:
        st.subheader('Uploaded Documents')
        selected_doc = st.selectbox('Select document', uploaded_docs)
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button('Preview'):
                doc_path = os.path.join(UPLOAD_DIR, selected_doc)
                content = read_file_from_disk(doc_path)
                st.session_state.preview_content = content[:2000]
        
        with col2:
            if st.button('Delete'):
                delete_uploaded_document(selected_doc)
                st.rerun()
        
        if 'preview_content' in st.session_state:
            with st.expander('Document Preview'):
                st.text(st.session_state.preview_content)
    
    st.header('Actions')
    if st.button('Rebuild Index'):
        with st.spinner('Rebuilding...'):
            rebuild_index()
        st.success('Index rebuilt')
        st.rerun()
    
    if st.button('Clear Uploaded Knowledge'):
        clear_uploaded_knowledge()
        st.success('Cleared')
        st.rerun()
    
    if st.button('Reset Chat'):
        st.session_state.messages = []
        st.rerun()
    
    st.header('Export')
    export_format = st.selectbox('Format', ['markdown', 'json'])
    if st.button('Export Conversation'):
        if 'messages' in st.session_state and st.session_state.messages:
            exported = export_conversation(st.session_state.messages, export_format)
            ext = 'md' if export_format == 'markdown' else 'json'
            st.download_button('Download', exported, file_name=f'conversation_{datetime.now().strftime("%Y%m%d_%H%M%S")}.{ext}')
        else:
            st.warning('No conversation to export')

if 'messages' not in st.session_state:
    st.session_state.messages = []

for idx, msg in enumerate(st.session_state.messages):
    with st.chat_message(msg['role']):
        st.markdown(msg['content'])
        
        if msg['role'] == 'assistant' and idx == len(st.session_state.messages) - 1:
            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button('👍', key=f'up_{idx}'):
                    save_feedback(st.session_state.messages[idx-1]['content'], msg['content'], 'positive')
                    st.success('Feedback saved')
            with col2:
                if st.button('👎', key=f'down_{idx}'):
                    save_feedback(st.session_state.messages[idx-1]['content'], msg['content'], 'negative')
                    st.success('Feedback saved')

query = st.chat_input('Ask MapMindGPT...')
if query:
    st.session_state.messages.append({'role': 'user', 'content': query})
    with st.chat_message('user'):
        st.markdown(query)
    
    with st.spinner('Retrieving context...'):
        context = retrieve(query, top_k=num_sources, scope=search_scope, use_reranker=use_reranker)
    
    history = st.session_state.messages[:-1]
    
    with st.chat_message('assistant'):
        answer = ask_llm(
            query,
            context,
            history,
            selected_model,
            selected_mode,
            use_cache=use_cache
        )

        st.markdown(answer)

        with st.expander('Sources Used'):
            if context:
                for c in context:
                    st.markdown(f"**{c['source']}** · `{c['category']}`")
                    st.code(format_snippet(c['text'], length=320))
            else:
                st.info('No relevant sources found')

        col1, col2 = st.columns([1, 1])

        with col1:
            if st.button('👍', key=f'feedback_up_{len(st.session_state.messages)}'):
                save_feedback(query, answer, 'positive')
                st.success('Thanks for feedback')

        with col2:
            if st.button('👎', key=f'feedback_down_{len(st.session_state.messages)}'):
                save_feedback(query, answer, 'negative')
                st.success('Thanks for feedback')
        
        st.session_state.messages.append({'role': 'assistant', 'content': answer})
